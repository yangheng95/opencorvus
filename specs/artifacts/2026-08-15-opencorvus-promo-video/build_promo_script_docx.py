from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
SOURCE = ROOT / "specs" / "records" / "2026-08" / "2026-08-15-opencorvus-promotion-video-production-script.md"
OUTPUT = ROOT / "specs" / "artifacts" / "2026-08-15-opencorvus-promo-video" / "OpenCorvus-宣传视频完整制作脚本.docx"

NAVY = "17324D"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
BLUE_PALE = "E8EEF5"
GRAY_PALE = "F2F4F7"
INK = "222222"
MUTED = "606A73"
WHITE = "FFFFFF"
GOLD = "A97718"
RED = "9B1C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent=120) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, *, latin="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None, italic=None):
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_keep(paragraph, *, next_=False, together=False, break_before=False):
    p_pr = paragraph._p.get_or_add_pPr()
    for tag, enabled in (("keepNext", next_), ("keepLines", together), ("pageBreakBefore", break_before)):
        node = p_pr.find(qn(f"w:{tag}"))
        if enabled and node is None:
            node = OxmlElement(f"w:{tag}")
            p_pr.append(node)


def set_paragraph_border(paragraph, *, side="left", color=BLUE, size=18, space=10):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def shade_paragraph(paragraph, fill=GRAY_PALE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_begin, instr, fld_sep, value, fld_end):
        run._r.append(node)
    set_run_font(run, size=9, color=MUTED)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22
    normal.font.color.rgb = RGBColor.from_string(INK)

    for name, size, color, before, after in (
        ("Title", 29, NAVY, 0, 8),
        ("Subtitle", 13.5, MUTED, 0, 14),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, BLUE_DARK, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.08

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.22

    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.text = "OPEN CORVUS  /  PROMOTION VIDEO PRODUCTION SCRIPT"
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(2)
        set_paragraph_border(hp, side="bottom", color="D7DBE2", size=6, space=4)
        for run in hp.runs:
            set_run_font(run, size=8.5, bold=True, color=MUTED)
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        label = fp.add_run("OpenCorvus · 录制执行手册  |  ")
        set_run_font(label, size=8.5, color=MUTED)
        add_page_field(fp)


def add_metric_strip(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340], indent=120)
    set_repeat_table_header(table.rows[0])
    items = [("主版本", "约 13 分钟"), ("画幅", "16:9 桌面端"), ("语言", "中文普通话"), ("结构", "33 个镜头")]
    for cell, (label, value) in zip(table.rows[0].cells, items):
        set_cell_shading(cell, BLUE_PALE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        set_run_font(r, size=8.5, bold=True, color=BLUE_DARK)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=10.5, bold=True, color=NAVY)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("VIDEO PRODUCTION PLAYBOOK")
    set_run_font(r, size=10, bold=True, color=GOLD)
    title = doc.add_paragraph(style="Title")
    title.add_run("OpenCorvus 宣传视频\n完整制作脚本")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("逐镜头操作 · 页面口述词 · 占位数据 · 专家团协议图 · 录制与成片验收")
    add_metric_strip(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("核心表达")
    set_run_font(r, size=11, bold=True, color=BLUE)
    lead = doc.add_paragraph()
    lead.paragraph_format.left_indent = Inches(0.16)
    lead.paragraph_format.right_indent = Inches(0.16)
    lead.paragraph_format.space_after = Pt(10)
    shade_paragraph(lead, "F4F6F9")
    set_paragraph_border(lead, side="left", color=BLUE, size=22, space=12)
    r = lead.add_run("OpenCorvus 是面向长程任务的开源 Agent 工作台：让能力、责任、消息、工具、证据与人工决定保持可见。")
    set_run_font(r, size=12, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("事实边界")
    set_run_font(r, size=10.5, bold=True, color=RED)
    for text in (
        "Mission 协调多个 Task，但不是一个更大的超级 Agent。",
        "自进化先比较证据，再由用户显式授权晋升或恢复。",
        "无人值守运行只在本地或托管 OpenCorvus 运行时在线期间继续。",
        "官方 Provider 数据与本地请求账本只对账，不相加。",
    ):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)
    doc.add_page_break()


def add_squad_diagram(doc: Document) -> None:
    doc.add_heading("专家团定义与协议图（可编辑版）", level=1)
    intro = doc.add_paragraph("后期按从左到右的四层逐步点亮；图中每一列对应旁白的一段。")
    intro.paragraph_format.space_after = Pt(8)
    table = doc.add_table(rows=2, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340], indent=120)
    set_repeat_table_header(table.rows[0])
    headers = ["01 声明层", "02 冻结层", "03 运行层", "04 治理层"]
    bodies = [
        "Manifest\nIdentity / Version / Digest\nSelector\nNamed Agents\nWorkflow DAG\nSkills / Tools / MCP\nConfiguration schema",
        "Mission held Squad IDs\n↓\nTask creation freezes exact revision\n↓\nVisible immutable workflow decision\nbefore first domain dispatch\n\n安装 ≠ 启用",
        "Orchestrator projection\n↓\nStreaming messages & tools\n↓\nTyped Artifacts\n+ Host observations",
        "Permission / Human takeover\nAccept / Retry / Replan\nEvolution Campaign\nCompare / Promote / Restore\n\n证据不足不自动晋升",
    ]
    colors = [BLUE_DARK, BLUE, NAVY, GOLD]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, colors[idx])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(headers[idx])
        set_run_font(r, size=10, bold=True, color=WHITE)
    for idx, cell in enumerate(table.rows[1].cells):
        set_cell_shading(cell, "F6F8FA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line_idx, line in enumerate(bodies[idx].split("\n")):
            if line_idx:
                p.add_run().add_break()
            r = p.add_run(line)
            set_run_font(r, size=9.2, bold=line_idx == 0, color=INK)
    boundary = doc.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    boundary.paragraph_format.space_before = Pt(8)
    shade_paragraph(boundary, "FFF4D6")
    r = boundary.add_run("Task 生命周期内不静默换团   ·   安装不等于启用   ·   证据不足不自动晋升")
    set_run_font(r, size=10.2, bold=True, color="6E4E00")


def clean_inline(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def add_rich_text(paragraph, text: str, *, default_size=10.5):
    tokens = re.split(r"(\[[A-Z0-9_]+\]|`[^`]+`|\*\*[^*]+\*\*)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("[") and token.endswith("]"):
            run = paragraph.add_run(token)
            set_run_font(run, size=default_size, bold=True, color=RED)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, latin="Consolas", east_asia="Microsoft YaHei", size=default_size - 0.5, color=BLUE_DARK)
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, bold=True, color=INK)
        else:
            run = paragraph.add_run(token)
            set_run_font(run, size=default_size, color=INK)


def add_quote(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.22
    shade_paragraph(p, "F4F6F9")
    set_paragraph_border(p, side="left", color=BLUE, size=22, space=12)
    label = p.add_run("口述词  ")
    set_run_font(label, size=9.2, bold=True, color=BLUE)
    add_rich_text(p, text, default_size=10.8)
    set_keep(p, together=True)


def add_markdown_table(doc: Document, lines: list[str]):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    cols = len(rows[0])
    widths = {2: [2400, 6960], 3: [1700, 3600, 4060], 4: [1500, 2500, 2500, 2860]}.get(cols, [9360 // cols] * cols)
    widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_geometry(table, widths, indent=120)
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        for cidx, value in enumerate(row):
            cell = table.rows[ridx].cells[cidx]
            if ridx == 0:
                set_cell_shading(cell, BLUE_PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_rich_text(p, value, default_size=8.8 if cols >= 3 else 9.2)
            for run in p.runs:
                if ridx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(BLUE_DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def extract_delivery_markdown() -> list[str]:
    all_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(all_lines) if line == "## 拍摄前准备")
    end = next(i for i, line in enumerate(all_lines) if line == "## 实施与验证顺序")
    return all_lines[start:end]


def add_delivery_body(doc: Document):
    lines = extract_delivery_markdown()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                shade_paragraph(p, "F7F7F7")
                p.paragraph_format.left_indent = Inches(0.12)
                for idx, code_line in enumerate(code_lines):
                    if idx:
                        p.add_run().add_break()
                    r = p.add_run(code_line)
                    set_run_font(r, latin="Consolas", east_asia="Microsoft YaHei", size=7.7, color=INK)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        if stripped.startswith("### S13 "):
            add_squad_diagram(doc)
        if stripped.startswith("## 完整逐镜头脚本"):
            doc.add_page_break()
        if stripped.startswith("## 专家团协议图的 Mermaid"):
            doc.add_page_break()
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_rich_text(p, clean_inline(stripped[3:]), default_size=16)
            set_keep(p, next_=True)
        elif stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_rich_text(p, clean_inline(stripped[4:]), default_size=13)
            set_keep(p, next_=True)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) < 90:
            p = doc.add_paragraph(style="Heading 3")
            add_rich_text(p, stripped, default_size=11.5)
            set_keep(p, next_=True)
        elif stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            add_quote(doc, " ".join(quote_lines))
            continue
        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\. ", "", stripped))
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:])
        elif stripped:
            p = doc.add_paragraph()
            add_rich_text(p, stripped)
        i += 1


def add_recording_overview(doc: Document):
    doc.add_heading("制作总览", level=1)
    p = doc.add_paragraph()
    p.add_run("定位：").bold = True
    p.add_run("面向第一次了解 OpenCorvus 的开发者、技术负责人、独立创作者与 Agent 工作流构建者。")
    p = doc.add_paragraph()
    p.add_run("目的：").bold = True
    p.add_run("先建立“可控、可审查的 Agent 工作台”认知，再用真实界面证明单 Agent、专家团协作和长程运行能力。")
    p = doc.add_paragraph()
    p.add_run("交付方式：").bold = True
    p.add_run("每个镜头同时给出页面路径、录制前状态、具体动作、屏幕结果、逐页面口述词、字幕、转场和验收点。")
    doc.add_heading("章节时间预算", level=2)
    rows = [
        ("00:00-00:35", "定位与背景"),
        ("00:35-04:15", "Chat / Work、Browser / Computer 与三个基础 CASE"),
        ("04:15-06:12", "专家团定义、协议与网页游戏专家团"),
        ("06:12-08:05", "专家团集群协作与软件工程 Mission"),
        ("08:05-09:35", "自进化与调研专家团"),
        ("09:35-10:35", "Mission 长任务与 Autonomous 调度"),
        ("10:35-12:02", "Provider、看板、Channel 与 Token"),
        ("12:02-13:10", "原则、承诺与加入邀请"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=2)
    set_table_geometry(table, [2200, 7160])
    for idx, text in enumerate(("时间", "章节")):
        set_cell_shading(table.rows[0].cells[idx], BLUE_PALE)
        r = table.rows[0].cells[idx].paragraphs[0].add_run(text)
        set_run_font(r, size=9.2, bold=True, color=BLUE_DARK)
    set_repeat_table_header(table.rows[0])
    for ridx, (time, section) in enumerate(rows, 1):
        table.rows[ridx].cells[0].paragraphs[0].add_run(time)
        table.rows[ridx].cells[1].paragraphs[0].add_run(section)
    doc.add_paragraph()


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_recording_overview(doc)
    add_delivery_body(doc)
    props = doc.core_properties
    props.title = "OpenCorvus 宣传视频完整制作脚本"
    props.subject = "逐镜头录屏、口述与剪辑执行手册"
    props.author = "OpenCorvus"
    props.keywords = "OpenCorvus, promotion video, recording, Expert Squad, Mission"
    props.comments = "Generated from the repository-authoritative production script record."
    doc.save(OUTPUT)
    print("docx-created")


if __name__ == "__main__":
    main()
